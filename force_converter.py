#!/usr/bin/env python3
"""
Force File Converter - Converts remaining files with multiple methods
"""

import os
import sys
import subprocess
from pathlib import Path
import logging

# Setup logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('force_conversion.log'),
        logging.StreamHandler()
    ]
)

def install_package(package):
    """Install a single package"""
    try:
        result = subprocess.run([sys.executable, '-m', 'pip', 'install', package], 
                              capture_output=True, text=True)
        if result.returncode == 0:
            logging.info(f"✓ Installed {package}")
            return True
        else:
            logging.warning(f"Failed to install {package}: {result.stderr}")
            return False
    except Exception as e:
        logging.warning(f"Error installing {package}: {e}")
        return False

def convert_with_libreoffice(input_file, output_file):
    """Convert using LibreOffice"""
    try:
        # Check if LibreOffice is available
        libreoffice_paths = [
            '/Applications/LibreOffice.app/Contents/MacOS/soffice',
            '/usr/bin/libreoffice',
            '/usr/local/bin/libreoffice',
            'libreoffice'
        ]
        
        libreoffice_cmd = None
        for path in libreoffice_paths:
            try:
                if path == 'libreoffice':
                    result = subprocess.run(['which', 'libreoffice'], capture_output=True, text=True)
                    if result.returncode == 0:
                        libreoffice_cmd = 'libreoffice'
                        break
                elif os.path.exists(path):
                    libreoffice_cmd = path
                    break
            except:
                continue
        
        if not libreoffice_cmd:
            logging.warning(f"LibreOffice not found, skipping: {os.path.basename(input_file)}")
            return False
        
        output_dir = os.path.dirname(output_file)
        os.makedirs(output_dir, exist_ok=True)
        
        # Try LibreOffice conversion
        cmd = [libreoffice_cmd, '--headless', '--convert-to', 'pdf', '--outdir', output_dir, input_file]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=120)
        
        if result.returncode == 0:
            # Check if file was created
            base_name = os.path.splitext(os.path.basename(input_file))[0]
            generated_pdf = os.path.join(output_dir, f"{base_name}.pdf")
            
            if os.path.exists(generated_pdf):
                if generated_pdf != output_file:
                    os.rename(generated_pdf, output_file)
                return True
        
        logging.warning(f"LibreOffice failed for {input_file}: {result.stderr}")
        return False
    except Exception as e:
        logging.warning(f"LibreOffice error for {input_file}: {e}")
        return False

def convert_with_pandoc(input_file, output_file):
    """Convert using pandoc"""
    try:
        # Check if pandoc is available
        pandoc_paths = [
            '/usr/local/bin/pandoc',
            '/opt/homebrew/bin/pandoc',
            '/usr/bin/pandoc',
            'pandoc'
        ]
        
        pandoc_cmd = None
        for path in pandoc_paths:
            try:
                if path == 'pandoc':
                    result = subprocess.run(['which', 'pandoc'], capture_output=True, text=True)
                    if result.returncode == 0:
                        pandoc_cmd = 'pandoc'
                        break
                elif os.path.exists(path):
                    pandoc_cmd = path
                    break
            except:
                continue
        
        if not pandoc_cmd:
            logging.warning(f"Pandoc not found, skipping: {os.path.basename(input_file)}")
            return False
        
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        
        cmd = [pandoc_cmd, input_file, '-o', output_file]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        
        if result.returncode == 0 and os.path.exists(output_file):
            return True
        
        logging.warning(f"Pandoc failed for {input_file}: {result.stderr}")
        return False
    except Exception as e:
        logging.warning(f"Pandoc error for {input_file}: {e}")
        return False

def convert_docx_manual(input_file, output_file):
    """Manual DOCX conversion using python-docx and reportlab"""
    try:
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        
        doc = Document(input_file)
        pdf_doc = SimpleDocTemplate(output_file, pagesize=letter)
        styles = getSampleStyleSheet()
        elements = []
        
        # Add title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
        )
        
        filename = os.path.basename(input_file)
        elements.append(Paragraph(filename, title_style))
        elements.append(Spacer(1, 12))
        
        # Process paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Determine style based on paragraph
                if para.style.name.startswith('Heading'):
                    style = styles['Heading2']
                else:
                    style = styles['Normal']
                
                try:
                    elements.append(Paragraph(para.text, style))
                    elements.append(Spacer(1, 6))
                except Exception:
                    # Fallback for problematic text
                    clean_text = ''.join(c for c in para.text if ord(c) < 127)
                    elements.append(Paragraph(clean_text, styles['Normal']))
                    elements.append(Spacer(1, 6))
        
        # Process tables
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if len(cell_text) > 50:
                        cell_text = cell_text[:47] + "..."
                    row_data.append(cell_text)
                table_data.append(row_data)
            
            if table_data:
                try:
                    pdf_table = Table(table_data)
                    pdf_table.setStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 8),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black)
                    ])
                    elements.append(pdf_table)
                    elements.append(Spacer(1, 12))
                except Exception:
                    pass
        
        pdf_doc.build(elements)
        return True
        
    except Exception as e:
        logging.warning(f"Manual DOCX conversion failed for {input_file}: {e}")
        return False

def convert_xlsx_manual(input_file, output_file):
    """Manual XLSX conversion"""
    try:
        import openpyxl
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, PageBreak, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
        
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        
        wb = openpyxl.load_workbook(input_file, data_only=True)
        doc = SimpleDocTemplate(output_file, pagesize=landscape(A4))
        styles = getSampleStyleSheet()
        elements = []
        
        # Add title
        filename = os.path.basename(input_file)
        elements.append(Paragraph(filename, styles['Title']))
        elements.append(Paragraph(" ", styles['Normal']))
        
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            
            # Add sheet name
            elements.append(Paragraph(f"Sheet: {sheet_name}", styles['Heading1']))
            
            # Get data
            data = []
            max_rows = min(100, ws.max_row)  # Limit rows
            max_cols = min(20, ws.max_column)  # Limit columns
            
            for row in ws.iter_rows(min_row=1, max_row=max_rows, 
                                  min_col=1, max_col=max_cols, values_only=True):
                if any(cell is not None for cell in row):
                    row_data = []
                    for cell in row:
                        if cell is not None:
                            cell_str = str(cell)
                            if len(cell_str) > 30:
                                cell_str = cell_str[:27] + "..."
                            row_data.append(cell_str)
                        else:
                            row_data.append('')
                    data.append(row_data)
            
            if data:
                try:
                    table = Table(data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 7),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('GRID', (0, 0), (-1, -1), 0.5, colors.black)
                    ]))
                    elements.append(table)
                except Exception:
                    elements.append(Paragraph("Table too large to display", styles['Normal']))
            
            if len(wb.sheetnames) > 1:
                elements.append(PageBreak())
        
        doc.build(elements)
        return True
        
    except Exception as e:
        logging.warning(f"Manual XLSX conversion failed for {input_file}: {e}")
        return False

def convert_png_manual(input_file, output_file):
    """Manual PNG conversion"""
    try:
        from PIL import Image
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Image as RLImage
        
        os.makedirs(os.path.dirname(output_file), exist_ok=True)
        
        with Image.open(input_file) as img:
            if img.mode != 'RGB':
                img = img.convert('RGB')
            
            doc = SimpleDocTemplate(output_file, pagesize=letter)
            page_width, page_height = letter
            
            img_width, img_height = img.size
            scale = min((page_width * 0.8) / img_width, (page_height * 0.8) / img_height)
            new_width = img_width * scale
            new_height = img_height * scale
            
            rl_img = RLImage(input_file, width=new_width, height=new_height)
            doc.build([rl_img])
        
        return True
        
    except Exception as e:
        logging.warning(f"Manual PNG conversion failed for {input_file}: {e}")
        return False

def force_convert_file(input_file, output_file):
    """Try multiple conversion methods"""
    ext = Path(input_file).suffix.lower()
    
    methods = []
    
    if ext in ['.docx', '.doc']:
        methods = [
            ("LibreOffice", convert_with_libreoffice),
            ("Pandoc", convert_with_pandoc),
            ("Manual DOCX", convert_docx_manual)
        ]
    elif ext in ['.xlsx', '.xls']:
        methods = [
            ("LibreOffice", convert_with_libreoffice),
            ("Manual XLSX", convert_xlsx_manual)
        ]
    elif ext in ['.pptx', '.ppt']:
        methods = [
            ("LibreOffice", convert_with_libreoffice),
            ("Pandoc", convert_with_pandoc)
        ]
    elif ext == '.png':
        methods = [
            ("Manual PNG", convert_png_manual)
        ]
    
    for method_name, method_func in methods:
        try:
            logging.info(f"Trying {method_name} for {os.path.basename(input_file)}")
            if method_func(input_file, output_file):
                if os.path.exists(output_file) and os.path.getsize(output_file) > 0:
                    logging.info(f"✓ Success with {method_name}: {os.path.basename(input_file)}")
                    return True
        except Exception as e:
            logging.warning(f"Method {method_name} failed: {e}")
    
    logging.error(f"✗ All methods failed for: {os.path.basename(input_file)}")
    return False

def copy_existing_pdfs(source_dir, target_dir):
    """Mavjud PDF fayllarni ko'chirish"""
    copied = 0
    
    for root, dirs, files in os.walk(source_dir):
        for file in files:
            if file.lower().endswith('.pdf'):
                source_file = os.path.join(root, file)
                
                # Nisbiy yo'lni hisoblash
                rel_path = os.path.relpath(root, source_dir)
                if rel_path == '.':
                    target_folder = target_dir
                else:
                    target_folder = os.path.join(target_dir, rel_path)
                
                os.makedirs(target_folder, exist_ok=True)
                target_file = os.path.join(target_folder, file)
                
                # Agar fayl mavjud bo'lmasa yoki o'lchami farq qilsa, ko'chirish
                if not os.path.exists(target_file) or os.path.getsize(source_file) != os.path.getsize(target_file):
                    try:
                        import shutil
                        shutil.copy2(source_file, target_file)
                        logging.info(f"✓ PDF ko'chirildi: {file}")
                        copied += 1
                    except Exception as e:
                        logging.warning(f"PDF ko'chirishda xatolik {file}: {e}")
    
    return copied

def main():
    source_dir = "/Users/bro/PROJECTS/fileconvertor/boltfiles"
    target_dir = "/Users/bro/PROJECTS/fileconvertor/converted"
    
    logging.info("Starting force conversion...")
    
    # Avval mavjud PDF fayllarni ko'chirish
    logging.info("Mavjud PDF fayllarni ko'chirish...")
    copied_pdfs = copy_existing_pdfs(source_dir, target_dir)
    if copied_pdfs > 0:
        logging.info(f"✓ {copied_pdfs} ta PDF fayl ko'chirildi")
    
    # Install required packages
    packages = ['python-docx', 'openpyxl', 'reportlab', 'Pillow']
    for package in packages:
        install_package(package)
    
    # Find all files that need conversion
    supported_extensions = {'.docx', '.doc', '.xlsx', '.xls', '.pptx', '.ppt', '.png'}
    
    files_to_convert = []
    for root, dirs, files in os.walk(source_dir):
        for file in files:
            file_path = Path(root) / file
            if file_path.suffix.lower() in supported_extensions:
                
                # Calculate output path
                rel_path = os.path.relpath(root, source_dir)
                if rel_path == '.':
                    output_dir = target_dir
                else:
                    output_dir = os.path.join(target_dir, rel_path)
                
                output_filename = file_path.stem + '.pdf'
                output_path = os.path.join(output_dir, output_filename)
                
                # Check if already converted
                if not (os.path.exists(output_path) and os.path.getsize(output_path) > 0):
                    files_to_convert.append((str(file_path), output_path))
    
    logging.info(f"Found {len(files_to_convert)} files to convert")
    
    # Convert files
    converted = 0
    failed = 0
    
    for i, (input_file, output_file) in enumerate(files_to_convert, 1):
        logging.info(f"Processing {i}/{len(files_to_convert)}: {os.path.basename(input_file)}")
        
        if force_convert_file(input_file, output_file):
            converted += 1
        else:
            failed += 1
    
    logging.info(f"\nForce conversion completed!")
    logging.info(f"Converted: {converted}")
    logging.info(f"Failed: {failed}")
    logging.info(f"Total processed: {len(files_to_convert)}")

if __name__ == "__main__":
    main()